"""
JARVIS Action Engine
Allows JARVIS to control the PC and Search the Web.
"""
import os
import sys
import webbrowser
import pyautogui
# from AppOpener import open as app_open
# from AppOpener import close as app_close
from duckduckgo_search import DDGS
import ollama
import json

# Safety settings
pyautogui.FAILSAFE = True

class ActionEngine:
    def __init__(self):
        pass
        
    def execute(self, action_text: str):
        """
        Parses and executes an action command.
        """
        if ":" not in action_text:
            return "❌ Invalid command format."
            
        cmd, param = action_text.split(":", 1)
        cmd = cmd.strip().upper()
        param = param.strip()
        
        print(f"⚙️ Executing: {cmd} -> {param}")
        
        try:
            if cmd == "OPEN":
                return self.open_app(param)
            elif cmd == "CLOSE":
                return self.close_app(param)
            elif cmd == "SEARCH":
                return self.web_search(param)
            elif cmd == "TYPE":
                return self.type_text(param)
            elif cmd == "SYSTEM":
                return self.system_control(param)
            elif cmd == "DOC":
                # Expecting "Filename | Content" or just Content
                if "|" in param:
                    fname, content = param.split("|", 1)
                else:
                    fname = "Jarvis_Report"
                    content = param
                return self.create_word_doc(fname.strip(), content.strip())
            else:
                return f"❌ Unknown command: {cmd}"
        except Exception as e:
            return f"❌ Execution failed: {e}"

    def create_word_doc(self, title, content):
        print(f"📝 Creating Word Doc: {title}...")
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            
            # Save
            filename = f"{title.replace(' ', '_')}.docx"
            path = os.path.abspath(filename)
            doc.save(path)
            
            # Open
            os.startfile(path)
            return f"Created and Opened {filename}"
        except ImportError:
            return "❌ python-docx not installed."
        except Exception as e:
            return f"❌ Doc Error: {e}"

    def open_app(self, app_name):
        print(f"🚀 Launching {app_name}...")
        try:
            from AppOpener import open as app_open
            app_open(app_name, match_closest=True, output=False)
            return f"Opened {app_name}"
        except:
            os.system(f"start {app_name}")
            return f"Attempted to open {app_name}"

    def close_app(self, app_name):
        print(f"🛑 Closing {app_name}...")
        try:
            from AppOpener import close as app_close
            app_close(app_name, match_closest=True, output=False)
            return f"Closed {app_name}"
        except:
            return f"Could not close {app_name}"

    def web_search(self, query):
        print(f"🦆 Searching DuckDuckGo for: {query}")
        blocklist = ["reddit.com", "wikipedia.org", "quora.com", "twitter.com", "facebook.com", "instagram.com"]
        try:
            results = DDGS().text(query, max_results=10) # Fetch more to filter
            filtered = []
            for r in results:
                link = r.get('href', '').lower()
                if not any(bad in link for bad in blocklist):
                    filtered.append(r)
                if len(filtered) >= 3: break
            
            if not filtered:
                return "No credible sources found."

            summary = "\n".join([f"- {r['title']}: {r['body']}" for r in filtered])
            return f"Found these credible results:\n{summary}"
        except Exception as e:
            # Fallback to browser
            url = f"https://duckduckgo.com/?q={query}"
            webbrowser.open(url)
            webbrowser.open(url)
            return f"Opened search in browser (API error: {e})"

    def predict_outcome(self, query):
        from datetime import datetime
        today = datetime.now().strftime("%Y-%m-%d")
        print(f"🔮 Oracle Mode: Predicting '{query}' (Date: {today})...")
        
        # 1. Multi-Angle Search (With Date for Freshness)
        sub_queries = [
            f"{query} price and statistics {today}",
            f"{query} latest news {today}",
            f"{query} expert forecast {today}"
        ]
        
        context = []
        for q in sub_queries:
            print(f"      🔎 Gathering data: {q}...")
            try:
                # Get top 3 results per sub-query to increase hit rate
                results = DDGS().text(q, max_results=3)
                for r in results:
                    context.append(f"Source: {r['title']} - {r['body']}")
            except:
                pass
                
        if not context:
            return "❌ Insufficient data to make a prediction."
            
        # 2. LLM Synthesis
        print("      🧠 Analyzing probability curves...")
        prompt = f"""
        You are JARVIS (The Oracle). British. Precise.
        TODAY'S DATE: {today}
        
        Task: Predict the outcome for Sir: "{query}" based *STRICTLY* on the data below.
        CRITICAL: Do NOT use internal training figures. Use the provided snippets.
        
        Data:
        {json.dumps(context)}
        
        Output format (Strict JSON):
        {{
            "prediction": "Short, British-style outcome (e.g. 'It appears highly likely, Sir...')",
            "confidence_percentage": "XX%",
            "key_factors": "List 3 main reasons"
        }}
        """
        
        try:
            res = ollama.chat(model="mixtral", messages=[{'role': 'user', 'content': prompt}])
            content = res['message']['content']
            # Try to parse JSON (sometimes models add chatter)
            if "{" in content and "}" in content:
                # Extract JSON part
                json_str = content[content.find("{"):content.rfind("}")+1]
                data = json.loads(json_str)
                return f"""
🔮 PREDICTION: {data['prediction']}
📊 PROBABILITY: {data['confidence_percentage']}
🔑 FACTORS: {data['key_factors']}
"""
            else:
                return f"🔮 Analysis:\n{content}"
        except Exception as e:
            return f"❌ Prediction Error: {e}"

    def type_text(self, text):
        print(f"⌨️ Typing: {text}")
        pyautogui.write(text, interval=0.05)
        return "Typed text"

    def system_control(self, command):
        cmd = command.lower()
        if "volume up" in cmd:
            pyautogui.press("volumeup", presses=5)
            return "Increased volume"
        elif "volume down" in cmd:
            pyautogui.press("volumedown", presses=5)
            return "Decreased volume"
        elif "mute" in cmd:
            pyautogui.press("volumemute")
            return "Muted audio"
        elif "lock" in cmd:
            os.system("rundll32.exe user32.dll,LockWorkStation")
            return "Locked workstation"
        else:
            return "Unknown system command"
